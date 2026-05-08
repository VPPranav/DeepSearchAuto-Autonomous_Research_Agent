import io

def generate_pdf(report: dict) -> io.BytesIO:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
    from reportlab.lib import colors
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
    styles = getSampleStyleSheet()
    
    # Custom Styles
    title_style = ParagraphStyle('CustomTitle', parent=styles['Title'], fontSize=24, textColor=colors.HexColor("#1e3a8a"), spaceAfter=20)
    h1_style = ParagraphStyle('Heading1', parent=styles['Heading1'], fontSize=16, textColor=colors.HexColor("#2563eb"), spaceAfter=10)
    h2_style = ParagraphStyle('Heading2', parent=styles['Heading2'], fontSize=14, textColor=colors.HexColor("#3b82f6"), spaceAfter=8)
    normal_style = styles['Normal']
    
    Story = []
    
    # Title
    Story.append(Paragraph(f"Research Report: {report.get('query', '')}", title_style))
    Story.append(Spacer(1, 12))
    
    # Executive Summary
    Story.append(Paragraph("Executive Summary", h1_style))
    Story.append(Paragraph(report.get('executive_summary', ''), normal_style))
    Story.append(Spacer(1, 12))
    
    # Key Findings
    Story.append(Paragraph("Key Findings", h1_style))
    Story.append(Spacer(1, 12))
    
    for i, topic in enumerate(report.get('sub_topics', [])):
        Story.append(Paragraph(f"{i+1}. {topic.get('question', '')}", h2_style))
        Story.append(Paragraph(topic.get('answer', ''), normal_style))
        Story.append(Spacer(1, 6))
        
        # Sources
        sources = topic.get('sources', [])
        if sources:
            Story.append(Paragraph("Sources:", ParagraphStyle('Sources', parent=normal_style, fontName='Helvetica-Bold')))
            for src in sources:
                Story.append(Paragraph(f"- <a href='{src.get('url', '')}'>{src.get('title', src.get('url', ''))}</a>", normal_style))
        Story.append(Spacer(1, 12))
        
    # Conclusion
    Story.append(Paragraph("Conclusion", h1_style))
    Story.append(Paragraph(report.get('conclusion', ''), normal_style))
    Story.append(Spacer(1, 12))
    
    # Confidence Score
    confidence = report.get('overall_confidence', 0)
    Story.append(Paragraph(f"Overall Confidence Score: {int(confidence * 100)}%", h2_style))
    
    doc.build(Story)
    buffer.seek(0)
    return buffer

def generate_docx(report: dict) -> io.BytesIO:
    try:
        from docx import Document
    except ImportError:
        pass
        
    doc = Document()
    doc.add_heading(f"Research Report: {report.get('query', '')}", 0)
    
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(report.get('executive_summary', ''))
    
    doc.add_heading('Key Findings', level=1)
    
    for i, topic in enumerate(report.get('sub_topics', [])):
        doc.add_heading(f"{i+1}. {topic.get('question', '')}", level=2)
        doc.add_paragraph(topic.get('answer', ''))
        
        sources = topic.get('sources', [])
        if sources:
            doc.add_heading('Sources:', level=3)
            for src in sources:
                doc.add_paragraph(f"{src.get('title', src.get('url', ''))} - {src.get('url', '')}", style='List Bullet')
                
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(report.get('conclusion', ''))
    
    confidence = report.get('overall_confidence', 0)
    doc.add_paragraph(f"Overall Confidence Score: {int(confidence * 100)}%", style='Intense Quote')
    
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
